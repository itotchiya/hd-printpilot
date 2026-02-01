"""
Extract text from Word document and save to markdown for comparison
"""
from docx import Document
import os

# Path to the docx file
docx_path = r"c:\Users\lenovo\Desktop\HD-PrintPilot\hd-printpilot\src\Docs\Excel Logic Extraction.docx"
output_path = r"c:\Users\lenovo\Desktop\HD-PrintPilot\hd-printpilot\src\Docs\docx-extracted.md"

# Load the document
doc = Document(docx_path)

# Extract all text
full_text = []

for para in doc.paragraphs:
    full_text.append(para.text)

# Also extract tables if any
for table in doc.tables:
    for row in table.rows:
        row_text = []
        for cell in row.cells:
            row_text.append(cell.text)
        full_text.append(" | ".join(row_text))

# Write to output file
with open(output_path, 'w', encoding='utf-8') as f:
    f.write('\n'.join(full_text))

print(f"Extracted {len(doc.paragraphs)} paragraphs and {len(doc.tables)} tables")
print(f"Saved to: {output_path}")
print("\n--- First 5000 characters ---\n")
content = '\n'.join(full_text)
print(content[:5000])
